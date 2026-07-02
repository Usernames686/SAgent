from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("/Users/chenwankang/Documents/SAgent/SAgent-master/SAgent学习理论与算法实现详解.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(98, 111, 128)
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
CALLOUT = "F4F6F9"
WHITE = RGBColor(255, 255, 255)
BLACK = RGBColor(0, 0, 0)


def set_run_font(run, size=None, color=None, bold=None, italic=None, east_asia="Microsoft YaHei", latin="Calibri"):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:ascii"), latin)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_para_spacing(paragraph, before=0, after=6, line=1.10):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def paragraph_border_bottom(paragraph, color="2E74B5", size="8", space="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.10


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_spacing(p, after=0)
    r = p.add_run("SAgent 学习理论与算法实现详解")
    set_run_font(r, size=9, color=MUTED)


def add_cover(doc):
    p = doc.add_paragraph()
    set_para_spacing(p, before=20, after=4)
    r = p.add_run("技术说明书")
    set_run_font(r, size=11, color=MUTED, bold=True)

    title = doc.add_paragraph()
    set_para_spacing(title, before=0, after=6)
    r = title.add_run("SAgent 项目的学习理论与算法实现详解")
    set_run_font(r, size=24, color=INK, bold=True)

    sub = doc.add_paragraph()
    set_para_spacing(sub, before=0, after=16)
    r = sub.add_run("自适应学习、知识追踪、能力评估、复习调度与 Agent 策略进化的系统化解读")
    set_run_font(r, size=13, color=MUTED)

    rows = [
        ("项目路径", "/Users/chenwankang/Documents/SAgent/SAgent-master"),
        ("文档范围", "后端学习算法、Agent 进化机制、理论文档与实现代码的对应关系"),
        ("生成日期", datetime.now().strftime("%Y-%m-%d")),
        ("阅读对象", "项目负责人、研发同学、算法评审、教学产品设计人员"),
    ]
    for label, value in rows:
        p = doc.add_paragraph()
        set_para_spacing(p, after=2)
        lr = p.add_run(f"{label}: ")
        set_run_font(lr, size=11, color=BLACK, bold=True)
        vr = p.add_run(value)
        set_run_font(vr, size=11, color=BLACK)

    rule = doc.add_paragraph()
    set_para_spacing(rule, before=12, after=12)
    paragraph_border_bottom(rule)

    callout = doc.add_table(rows=1, cols=1)
    set_table_geometry(callout, [9360])
    cell = callout.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    set_para_spacing(p, after=2)
    r = p.add_run("核心结论")
    set_run_font(r, size=11, color=INK, bold=True)
    p = cell.add_paragraph()
    set_para_spacing(p, after=0)
    r = p.add_run(
        "SAgent 的学习理论不是单一模型，而是一套组合式自适应学习框架："
        "BKT 追踪知识点掌握度，IRT 估计学生能力和题目难度匹配，"
        "强化学习/多臂老虎机思想优化路径推荐，SM-2 管理长期复习，"
        "A/B 测试和灰度发布机制让教学策略持续进化。"
    )
    set_run_font(r, size=11, color=BLACK)

    doc.add_page_break()


def add_p(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    set_para_spacing(p)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_para_spacing(p, after=4, line=1.167)
        r = p.add_run(item)
        set_run_font(r)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_para_spacing(p, after=8, line=1.167)
        r = p.add_run(item)
        set_run_font(r)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    repeat_table_header(table.rows[0])
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], header_fill)
        p = hdr[i].paragraphs[0]
        set_para_spacing(p, after=0)
        r = p.add_run(h)
        set_run_font(r, size=10.5, color=INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            set_para_spacing(p, after=0, line=1.10)
            r = p.add_run(str(val))
            set_run_font(r, size=10)
    set_table_geometry(table, widths)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    style_doc(doc)
    add_footer(doc.sections[0])
    add_cover(doc)

    doc.add_heading("1. 总体判断：这个项目到底在“学习”什么", level=1)
    add_p(doc, "从代码和文档看，SAgent 的“学习”有两层含义。第一层是学生学习：系统持续估计学生对知识点的掌握程度、能力水平、遗忘风险和适合的学习节奏。第二层是 Agent 自身进化：系统对不同教学策略、路径算法和 Prompt 策略做实验，依据学习效果逐步放大更好的策略。")
    add_p(doc, "所以它不是在训练一个端到端神经网络，也不是简单的题库系统，而是一个面向编程学习场景的自适应学习系统。其核心是把教育测量、认知记忆、推荐系统、强化学习思想和工程化 A/B 测试组合起来。")

    add_table(
        doc,
        ["层次", "关注对象", "核心问题", "项目中的主要方法"],
        [
            ("学生模型", "单个学习者", "学生会什么、不会什么、多久会忘", "BKT、IRT、知识衰减、SM-2"),
            ("路径决策", "下一个学习动作", "下一步学什么、用什么模式、难度多大", "BKT 收益、IRT 难度匹配、UCB 探索、前置知识约束"),
            ("教学策略", "Agent 的提示方式和教学风格", "哪种教法更有效", "A/B 测试、灰度发布、效果量、回滚"),
            ("长期优化", "系统整体策略", "如何根据真实数据变得更好", "策略变体、实验日志、人工审核、参数更新"),
        ],
        [1500, 1700, 2700, 3460],
        BLUE_GRAY,
    )

    doc.add_heading("2. 总体闭环：数据如何变成个性化决策", level=1)
    add_p(doc, "SAgent 的自适应闭环可以理解为五步：采集学习行为、更新学生状态、选择学习动作、记录结果、优化策略。这个闭环既作用于单个学生，也作用于整个 Agent 策略池。")
    add_numbers(doc, [
        "采集行为数据：包括答题是否正确、得分、耗时、提交状态、错误次数、学习节点、复习结果等。",
        "更新学习者状态：用 BKT 更新知识点掌握概率，用 IRT 更新能力值 theta，用指数衰减或 SM-2 处理遗忘。",
        "做出自适应决策：筛选前置知识已满足且尚未掌握的知识点，再结合能力、难度、连续答对/答错和画像选择下一步。",
        "评估学习效果：把结果转化为掌握度、正确率、通过率、挫败水平、用时等指标。",
        "进化教学策略：对策略变体做 A/B 实验，显著有效则逐步放量，无效或负面则回滚。",
    ])

    doc.add_heading("3. BKT：贝叶斯知识追踪", level=1)
    add_p(doc, "BKT 是项目中最核心的知识掌握模型。它把每个知识点看成一个隐藏状态：学生是否已经掌握。系统看不到真实掌握状态，只能通过答题、练习和提交结果进行概率推断。")
    add_table(
        doc,
        ["参数", "含义", "项目默认值/示例", "解释"],
        [
            ("P(L0)", "初始掌握概率", "0.15", "新知识点默认认为学生掌握概率较低"),
            ("P(T)", "学习转移概率", "0.20", "一次学习后从未掌握转为掌握的概率"),
            ("P(G)", "猜测概率", "0.15", "没掌握但答对的概率"),
            ("P(S)", "失误概率", "0.10", "掌握了但答错的概率"),
            ("P(F)", "遗忘概率", "0.05", "随时间遗忘导致掌握度下降的概率"),
        ],
        [1100, 1900, 1700, 4660],
    )
    add_p(doc, "项目中 BKT 的更新逻辑是：学生作答后，先考虑本次学习带来的转移概率，再根据答对或答错做贝叶斯后验更新，最后将概率限制在合理区间，避免出现 0 或 1 的极端值。")
    add_p(doc, "答对时，系统认为“学生掌握”的概率上升，但会考虑猜对可能；答错时，系统认为掌握概率下降，但会考虑失误可能。这比简单地用正确率更稳健，因为它区分了“会但粗心”和“不会但猜对”。")
    add_p(doc, "代码位置：apps/api/src/modules/agent/evolution/bkt-rl-algorithm.ts 与 apps/api/src/modules/vibe-learning/adaptive-learning.engine.ts。")

    doc.add_heading("4. IRT：项目反应理论与能力诊断", level=1)
    add_p(doc, "IRT 用来估计学生的整体能力值 theta，并选择合适难度的题目。它解决的问题不是“某个知识点会不会”，而是“这个学生当前能力大概在哪个水平”。")
    add_p(doc, "项目文档采用三参数 Logistic 模型：P(正确) = c + (1-c) / (1 + exp(-a(theta - b)))。其中 theta 是学生能力，a 是题目区分度，b 是题目难度，c 是猜测率。")
    add_table(
        doc,
        ["符号", "名称", "作用"],
        [
            ("theta", "能力值", "表示学生能力水平，范围通常按 -3 到 3 理解"),
            ("a", "区分度", "题目对不同能力学生的区分能力，越高越敏感"),
            ("b", "难度", "题目难度，越高表示越难"),
            ("c", "猜测率", "选择题中低能力学生靠猜答对的概率"),
        ],
        [1200, 1800, 6360],
    )
    add_p(doc, "在实现上，诊断服务会根据学生答题序列估计 theta、标准误、置信区间，并通过 Fisher 信息量选择下一道更有诊断价值的题。自适应学习引擎中还用简化版 IRT 来预测某个难度下的正确率，并推荐接近学生能力的练习难度。")
    add_p(doc, "代码位置：apps/api/src/modules/assessment/irt-assessment.service.ts 与 apps/api/src/modules/vibe-learning/adaptive-learning.engine.ts。")

    doc.add_heading("5. 路径推荐：BKT + IRT + UCB 的混合决策", level=1)
    add_p(doc, "路径推荐不是单纯按课程目录顺序推进，而是综合多种信号打分。项目中的推荐分数大致由四类因素构成：知识点掌握缺口、历史学习收益、探索奖励、当前阶段连贯性。")
    add_table(
        doc,
        ["因素", "对应理论", "项目中的作用"],
        [
            ("掌握缺口", "BKT", "掌握度低且接近可学习区间的知识点更值得推荐"),
            ("难度匹配", "IRT", "题目/节点难度应接近学生能力，避免过难或过易"),
            ("探索奖励", "UCB 多臂老虎机", "对推荐次数少但可能有效的知识点保留探索机会"),
            ("连续表现", "自适应节奏控制", "连续答对可加速，连续答错则降难度或回阅读模式"),
            ("前置约束", "知识图谱/DAG", "前置知识未满足时不推荐后续知识点"),
        ],
        [1700, 1900, 5760],
        BLUE_GRAY,
    )
    add_p(doc, "这里的强化学习更准确地说是“轻量级强化学习思想”或“上下文推荐启发式”，不是 PPO、DQN 这类完整深度 RL。代码里有 Q 值、奖励、UCB 探索项，但 Q 值更新是基于历史经验的简化平均回报。")
    add_p(doc, "代码位置：apps/api/src/modules/agent/evolution/bkt-rl-algorithm.ts。")

    doc.add_heading("6. SM-2：间隔重复与长期记忆", level=1)
    add_p(doc, "学习系统不能只关心当下答对，还要安排何时复习。项目使用 SM-2 算法，即 SuperMemo 经典间隔重复算法，根据回忆质量决定下次复习间隔。")
    add_p(doc, "项目把 0-100 分数映射为 0-5 的 quality：95 分以上为 5，85 分以上为 4，70 分以上为 3；低于 3 表示回忆失败，需要把间隔重置为 1 天。")
    add_table(
        doc,
        ["回忆质量", "含义", "调度结果"],
        [
            ("0-2", "回忆失败或困难", "间隔重置为 1 天，状态可能降为 learning"),
            ("3", "勉强通过", "进入复习链路，但难度因子增长较慢"),
            ("4", "较好回忆", "维持通过/掌握状态，间隔扩大"),
            ("5", "完美回忆", "难度因子更高，复习间隔增长更快"),
        ],
        [1600, 3200, 4560],
    )
    add_p(doc, "代码位置：apps/api/src/modules/vibe-learning/spaced-repetition.service.ts。")

    doc.add_heading("7. Agent 进化：A/B 测试、灰度发布与回滚", level=1)
    add_p(doc, "SAgent 的 Evolution Agent 不是神秘的自我意识式进化，而是工程化策略优化。它把教学策略、路径参数或 Prompt 变体作为实验对象，用真实学习数据比较效果。")
    add_numbers(doc, [
        "创建策略变体：例如 conservative、balanced、aggressive 三类 BKT 参数。",
        "启动 A/B 实验：设置对照组和实验组，从小流量开始。",
        "采集指标：样本量、完成率、平均得分、通过率、错误率、耗时等。",
        "统计判断：计算效果量和显著性。",
        "灰度推进：效果达到阈值后逐步扩大流量。",
        "人工审核：重要阶段需要人工确认。",
        "负向回滚：实验组效果变差时自动回到稳定版本。",
    ])
    add_p(doc, "这套机制的理论来源包括在线实验、统计检验、灰度发布和进化策略。项目文档把它类比为自然选择：变异是策略变体，选择是效果评估，遗传是保留好策略，适应是持续调整。")
    add_p(doc, "代码位置：apps/api/src/modules/agent/evolution/evolution-engine.service.ts、ab-test.service.ts、evolution-data-collector.service.ts。")

    doc.add_heading("8. 学习者画像与节奏控制", level=1)
    add_p(doc, "项目把学生分为 beginner、transition、advanced 三类画像，并为不同画像设置不同的 BKT 参数和初始 theta。这样做的意义是避免对所有学生使用同一套判断标准。")
    add_table(
        doc,
        ["画像", "初始能力/参数倾向", "学习策略含义"],
        [
            ("beginner", "低先验、低转移、高遗忘/高失误容忍", "更慢、更基础、更多提示和阅读"),
            ("transition", "中等初始能力", "允许跳过已掌握基础，聚焦新技术栈"),
            ("advanced", "高先验、高转移、低失误", "更快进入挑战、编码和项目实战"),
        ],
        [1700, 3300, 4360],
    )
    add_p(doc, "节奏控制则通过连续答对、连续答错和 paceScore 实现。连续答对时系统可推荐 coding 或更高难度；连续答错时则降回 reading，降低挫败感。")

    doc.add_heading("9. 知识图谱与前置依赖", level=1)
    add_p(doc, "项目中知识点不是孤立节点，而是有前置、依赖和阶段关系。路径推荐时会检查 prerequisites：如果前置知识掌握度不足，就不会推荐后续节点。")
    add_p(doc, "这相当于用知识图谱或 DAG 约束推荐系统，避免学生跳到尚不具备基础的内容上。文档中还提到 LightGCN、知识感知协同过滤等方向，但当前代码更偏规则化图约束和路径状态机。")

    doc.add_heading("10. 理论文档中的未来方向与当前实现差异", level=1)
    add_p(doc, "doc/理论.md 中列出了很多算法，其中一部分已经实现，一部分更像路线图或建议。理解这个差异很重要，否则容易误以为项目已经具备完整深度学习推荐系统。")
    add_table(
        doc,
        ["理论/算法", "当前状态", "说明"],
        [
            ("BKT", "已实现", "掌握概率更新、画像参数、遗忘衰减均有代码"),
            ("IRT", "已实现", "诊断题池、theta 估计、信息量选题、难度匹配"),
            ("SM-2", "已实现", "复习质量映射、间隔和 easeFactor 更新"),
            ("UCB/Q 值推荐", "部分实现", "有轻量经验池和 UCB 探索，不是完整 RL 框架"),
            ("A/B 测试进化", "已实现", "实验、灰度、效果量、回滚、审核链路"),
            ("DKT/SAKT", "文档建议", "尚未看到 LSTM/Transformer 知识追踪训练代码"),
            ("LightGCN/协同过滤", "文档建议", "当前更偏知识图谱约束与规则推荐"),
            ("FSRS", "文档建议", "当前实际为 SM-2"),
            ("HMM 情绪识别", "文档建议", "目前更偏行为规则和画像推断思路"),
            ("DSPy/APE Prompt 优化", "文档建议", "当前主要是人工策略变体 + A/B 实验"),
        ],
        [1900, 1600, 5860],
        BLUE_GRAY,
    )

    doc.add_heading("11. 关键代码地图", level=1)
    add_table(
        doc,
        ["模块", "文件", "作用"],
        [
            ("自适应学习", "apps/api/src/modules/vibe-learning/adaptive-learning.engine.ts", "BKT、IRT、画像、节奏、复习衰减和下一步决策"),
            ("BKT + RL 推荐", "apps/api/src/modules/agent/evolution/bkt-rl-algorithm.ts", "掌握度更新、经验池、Q 值、UCB 推荐分"),
            ("间隔重复", "apps/api/src/modules/vibe-learning/spaced-repetition.service.ts", "SM-2 复习调度"),
            ("IRT 诊断", "apps/api/src/modules/assessment/irt-assessment.service.ts", "能力估计、试题信息量、下一题选择"),
            ("进化引擎", "apps/api/src/modules/agent/evolution/evolution-engine.service.ts", "策略变体、实验启动、推进、审核和日志"),
            ("A/B 测试", "apps/api/src/modules/agent/evolution/ab-test.service.ts", "实验指标、显著性、效果量、灰度和回滚"),
            ("数据采集", "apps/api/src/modules/agent/evolution/evolution-data-collector.service.ts", "汇总提交、通过率、错误率并触发实验建议"),
            ("理论文档", "doc/理论.md", "算法理论总览和未来建议"),
            ("进化说明", "doc/Evolution Agent 的进化原理.md", "用案例解释策略进化流程"),
        ],
        [1600, 4300, 3460],
    )

    doc.add_heading("12. 用一句业务语言解释项目价值", level=1)
    add_p(doc, "SAgent 的价值在于：它试图把“每个学生应该学什么、什么时候复习、题目该多难、Agent 应该怎么教”这些问题，从固定规则变成数据驱动的动态决策。")
    add_p(doc, "对学生来说，系统希望做到不过早挑战、不重复已会内容、不让遗忘长期积累；对教学 Agent 来说，系统希望通过真实学习效果识别更好的教学策略，而不是凭主观感觉长期固定一套 Prompt。")

    doc.add_page_break()
    doc.add_heading("13. 评审建议", level=1)
    add_bullets(doc, [
        "如果要对外宣称“自进化”，建议明确说明当前是基于 A/B 测试、灰度发布和参数策略变体的工程化进化，不是自动训练大模型。",
        "如果要提高算法可信度，应补充离线评估集、实验样本量要求、显著性阈值说明和防止数据稀疏误判的保护。",
        "如果要进一步升级个性化，可以优先做两件事：真实学生数据上的 BKT 参数拟合，以及更稳定的上下文 bandit 推荐。",
        "如果要升级长期记忆调度，可以把 SM-2 与真实复习结果打通，后续再评估是否引入 FSRS。",
        "引入 DKT/SAKT 前，应先确认数据量足够；否则当前阶段 BKT 更适合。",
    ])

    doc.add_heading("附录：核心理论速查", level=1)
    add_table(
        doc,
        ["理论", "解决的问题", "一句话理解"],
        [
            ("BKT", "知识点掌握度", "用贝叶斯方法根据答题证据更新“会不会”的概率"),
            ("IRT", "能力与题目难度", "用题目参数反推学生能力，并选择信息量最大的题"),
            ("UCB/MAB", "探索与利用", "既推荐看起来好的，也保留尝试新策略的机会"),
            ("SM-2", "复习间隔", "记得越牢，下次复习越晚；忘得越多，越快复习"),
            ("A/B 测试", "策略效果验证", "让真实数据决定哪种教学策略更有效"),
            ("知识图谱", "学习顺序约束", "前置知识没掌握，就不贸然进入后续知识"),
        ],
        [1500, 2500, 5360],
        BLUE_GRAY,
    )

    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
